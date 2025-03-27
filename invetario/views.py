from django.shortcuts import render,redirect,get_object_or_404
from django.contrib import messages
from .models import Departameto,Computadora,Propiedad,LectorCD_DVD,MemoriaRam,Discos
from .forms import DepartamentoForm,ComputadoraForm,PropiedadForm,LectorForm,DiscoForm,MemoriaRamForm
from django.http import HttpResponse
from django.views import View
from docx import Document
from docx.shared import Inches
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from io import BytesIO  # Para manejar el PDF en memoria
from django.db.models import Q
from .paginacion import paginacion,paga_Size


class ComputadorasPDFView(View):
    def get(self, request, *args, **kwargs):
        # 1. Preparar la respuesta HTTP
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="computadora_list.pdf"'
        
        # 2. Crear un buffer en memoria para el PDF
        buffer = BytesIO()
        
        # 3. Crear el objeto PDF usando SimpleDocTemplate
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        elements = []  # Lista para almacenar los elementos del PDF

        # 4. Obtener los datos de las computadoras
        computadoras = Computadora.objects.order_by("numero_pc")
        
        #5. Preparar los datos para la tabla
        # Encabezados
        data = [['Componente Pc', '# Inventario', 'Marca', 'Estado','#_pc','Departamento']]  
        for computadora in computadoras:
            data.append([
                computadora.componente_pc,
                computadora.num_inventario,
                computadora.marca,
                computadora.estado,
                computadora.numero_pc,
                computadora.departamento
               ])
        # 6. Crear la tabla
        table = Table(data)
        # 7. Estilo de la tabla (opcional)
        style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.gray),  # Encabezado gris
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),  # Filas alternas
            ('GRID', (0, 0), (-1, -1), 1, colors.black)  # Bordes
        ])
        table.setStyle(style)
        # 8. Añadir la tabla a los elementos del documento
        elements.append(table)

        # 9. Construir el PDF
        doc.build(elements)
        
        # 10. Obtener el valor del buffer y escribirlo en la respuesta
        pdf = buffer.getvalue()
        buffer.close()
        response.write(pdf)

        return response

class ComputadorasListWordView(View):
       def get(self, request, *args, **kwargs):
            # 1. Preparar la respuesta HTTP
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename="task_list.docx"'

            # 2. Crear un buffer en memoria
            buffer = BytesIO()

            # 3. Crear un nuevo documento Word
            document = Document()
            
            # 4. Agregar un título (opcional)
            document.add_heading('Lista de Componentes de la Pc', 0)
            # 5. Obtener los datos de las Componentes Pc
            computadoras = Computadora.objects.order_by("numero_pc")
            
            # 6. Agregar la tabla
            table = document.add_table(rows=1, cols=6)  # 6 columnas:Componentes de Pc, # Inev,Departamento , 
            
            # 7. Encabezados de la tabla
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Componentes de Pc'
            hdr_cells[1].text = ' Número de Inventario'
            hdr_cells[2].text = 'Marca'
            hdr_cells[3].text = 'Estado'
            hdr_cells[4].text = 'Número de Pc'
            hdr_cells[5].text = 'Departamento'
            # 8. Agregar datos a la tabla
            for computadora in computadoras:
                row_cells = table.add_row().cells
                row_cells[0].text = computadora.componente_pc
                row_cells[1].text = str(computadora.num_inventario)
                row_cells[2].text = computadora.marca
                row_cells[3].text = computadora.estado
                row_cells[4].text = str(computadora.numero_pc)
                row_cells[5].text = computadora.departamento.nombre
               
            # 9. Guardar el documento en el buffer
            document.save(buffer)
            # 10. Obtener el valor del buffer y escribirlo en la respuesta
            word_file = buffer.getvalue()
            buffer.close()
            response.write(word_file)

            return response


#Listar Componentes pc  

def base(request):
    return render(request,'base.html')

def listar_computadoras(request):
    computadoras = Computadora.objects.order_by("departamento")
    # Obtener el valor de búsqueda de la URL
    search = request.GET.get('search', '')  # Si no hay valor, usa una cadena vacía
    # Filtrar las computadoras por departamento si hay un valor de búsqueda
    if search:
          computadoras = computadoras.filter(
            Q(departamento__nombre__icontains=search) | 
            Q(marca__icontains=search)
        )
    #Esto es para paginar 
    page_size = paga_Size(request)
    page_obj=paginacion(request,computadoras)
    context = {
        'page_obj': page_obj,
        'page_size': str(page_size),  # Pasar page_size como cadena al contexto
        'search': search,  # Pasar el valor de búsqueda al contexto
    }
    return render(request,'computadoras/listar.html',context)



def crear_computadora(request):
    if request.method == 'POST':
        form = ComputadoraForm(request.POST)
        numero_pc=int(request.POST['numero_pc'])
        if numero_pc < 0 :
            return render(request,'computadoras/crear.html',{
                'form':form,
                'error':'El numero de pc debe ser mayor que 0 no menor'
            })
        if form.is_valid():
            form.save()
            messages.success(request,'Componente de la pc guardado con éxito')
            return redirect('listar_computadoras')
    else:
        form = ComputadoraForm()
    return render(request,'computadoras/crear.html',{'form':form})

def editar_computadora(request,pk):
    computadora = get_object_or_404(Computadora,pk=pk)
    if request.method=="POST":
        form = ComputadoraForm(request.POST,instance=computadora)
        if form.is_valid():
            form.save()
            messages.success(request,'Computadora editada con éxito')
            return redirect('listar_computadoras')
        return render(request,'computadoras/editar.html',{'form':form})
    else:
        form = ComputadoraForm(instance=computadora)
    return render(request,'computadoras/editar.html',{'form':form})

def eliminar_computadora(request,pk):
    computadora = get_object_or_404(Computadora,pk=pk)
    if computadora:
        computadora.delete()
        messages.success(request,'Computadora eliminada con éxito')
        return redirect('listar_computadoras')
    
        
def listar_departamentos(request):
    departamentos = Departameto.objects.order_by("nombre")
    return render(request,'departamentos/listar.html',{'departamentos':departamentos})

def crear_departamento(request):
    if request.method == 'POST':
        form = DepartamentoForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request,'Departamento guardado con éxito')
            return redirect('listar_departamentos')
    else:
        form = DepartamentoForm()
    return render(request,'departamentos/crear.html',{'form':form})

def editar_departamento(request,pk):
    departamento = get_object_or_404(Departameto,pk=pk)
    if request.method == 'POST':
        form = DepartamentoForm(request.POST,instance=departamento)
        if form.is_valid():
            form.save()
            messages.success(request,'Departamento editado con éxito')
            return redirect('listar_departamentos')
        return render(request,'departamentos/editar.html',{'form':form})
    else:
        form = DepartamentoForm(instance=departamento)
    return render(request,'departamentos/editar.html',{'form':form})

def eliminar_departamento(request,pk):
    departamento = get_object_or_404(Departameto,pk=pk)
    if departamento:
        departamento.delete()
        messages.success(request,'Departamento eliminado con éxito')
        return redirect('listar_departamentos')
    

#listar memorias ram
def listar_memorias_ram(request):
    lista_memorias_ram=MemoriaRam.objects.all()
    search=request.GET.get('search','')
    if search:
        lista_memorias_ram=lista_memorias_ram.filter(Q(marca__icontains=search))
    return render(request,'memorias/listar.html',{'lista_memorias_ram':lista_memorias_ram})
#crear memoria ram

def crear_memoria_ram(request):
    if request.method=='POST':
        form=MemoriaRamForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request,'!Memoria Ram guardada con éxito')
            return redirect('listar_memorias_ram')
        return render(request,'memorias/crear.html',{'form':form})
    else:
        form = MemoriaRamForm()
    return render(request,'memorias/crear.html',{'form':form})

def editar_memoria_ram(request,pk):
    memoria_ram = get_object_or_404(MemoriaRam,pk=pk)
    if request.method == 'POST':
        form = MemoriaRamForm(request.POST,instance=memoria_ram)
        if form.is_valid():
            form.save()
            messages.success(request,'!Memoria Ram editad con éxito')
            return redirect('listar_memorias_ram')
        return render(request,'memorias/editar.html',{'form':form})
    else:
        form = MemoriaRamForm(instance=memoria_ram)
    return render(request,'memorias/editar.html',{'form':form})

def eliminar_memoria_ram(request,pk):
    memoria_ram = get_object_or_404(MemoriaRam,pk=pk)
    if memoria_ram:
        memoria_ram.delete()
        return redirect('listar_memorias_ram')

#función listado de discos duros
def listar_discos_duros(request):
    lista_discos_duros=Discos.objects.all()
    return render(request,'discos/listar.html',{'lista_discos_duros':lista_discos_duros})

#crear disco duro
def crear_disco_duro(request):
    if request.method=='POST':
        form = DiscoForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request,'!Disco duro guardado con éxito')
            return redirect('listar_discos_duros')
        return render(request,'discos/crear.html',{'form':form})
    else:
        form=DiscoForm()
    return render(request,'discos/crear.html',{'form':form})

def editar_disco_duro(request,pk):
    disco_duro=get_object_or_404(Discos,pk=pk)
    if request.method=="POST":
        form=DiscoForm(request.POST,instance=disco_duro)
        if form.is_valid():
            form.save()
            messages.success(request,'!Disco duro guardado con éxito')
            return redirect('listar_discos_duros')
        return render(request,'discos/crear.html',{'form':form})
    else:
        form=DiscoForm(instance=disco_duro)
    return render(request,'discos/crear.html',{'form':form})
            
def eliminar_disco_duro(request,pk):
    disco_duro=get_object_or_404(Discos,pk=pk)
    if disco_duro:
        disco_duro.delete()
        return redirect('listar_discos_duro')

#funcion de listar lectores cd-dvd
def listar_lectores(request):
    lista_lectores=LectorCD_DVD.objects.all()
    return render(request,'lectores/listar.html',{'lista_lectores':lista_lectores})

#crear lector 
def crear_lector(request):
    if request.method=='POST':
        form = LectorForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request,'!Lector guardado con éxito')
            return redirect('listar_lectores')
        return render(request,'lectores/crear.html',{'form':form})
    else:
        form=LectorForm()
    return render(request,'lectores/crear.html',{'form':form})


def editar_lector(request,pk):
    lector=get_object_or_404(LectorCD_DVD,pk=pk)
    if request.method=="POST":
        form=LectorForm(request.POST,instance=lector)
        if form.is_valid():
            form.save()
            messages.success(request,'! Lector guardado con éxito')
            return redirect('listar_lectores')
        return render(request,'lectores/crear.html',{'form':form})
    else:
        form=LectorForm(instance=lector)
    return render(request,'lectores/crear.html',{'form':form})
            
def eliminar_lector(request,pk):
    lector=get_object_or_404(LectorCD_DVD,pk=pk)
    if lector:
        lector.delete()
        return redirect('listar_lectores')

#listar Propiedades
def listar_propiedades(request):
    propiedades = Propiedad.objects.all()
    # Obtener el valor de búsqueda de la URL
    search = request.GET.get('search', '')  # Si no hay valor, usa una cadena vacía
    # Filtrar las computadoras por departamento si hay un valor de búsqueda
    if search:
        propiedades = propiedades.filter(Q(procesador__icontains=search))
    #Esto es para paginar 
    page_size = paga_Size(request)
    page_obj=paginacion(request,propiedades)
    context = {
        'page_obj': page_obj,
        'page_size': str(page_size),  # Pasar page_size como cadena al contexto
        'search': search,  # Pasar el valor de búsqueda al contexto
    }
    return render(request,'propiedades/listar.html',context)

def crear_propiedad(request):
    form=PropiedadForm(request.POST)
    computadoras = Computadora.objects.all()
    if form.is_valid():
            form.save()
            messages.success(request,'!Propiedades guardado con éxito')
            return redirect('listar_propiedades')
    else:
        form = PropiedadForm()
    return render(request,'propiedades/crear.html',{'form':form,'computadoras':computadoras})

def editar_propiedad(request,pk):
    propiedad=get_object_or_404(Propiedad,pk=pk)
    form=PropiedadForm(request.POST)
    if request.method == 'POST':
        form = PropiedadForm(request.POST,instance=propiedad)
        if form.is_valid():
            form.save()
            messages.success(request,'!Propiedad editada con éxito')
            return redirect('listar_propiedades')
        return render(request,'propiedades/editar.html',{'form':form})
    else:
        form = PropiedadForm(instance=propiedad)
    return render(request,'propiedades/editar.html',{'form':form})

def ver_propiedades(request,pk):
    propiedad=get_object_or_404(Propiedad,pk=pk)
    if propiedad:
        return render(request,'propiedades/ver.html',{'propiedad':propiedad})
    

def eliminar_propiedad(request,pk):
    propiedad=get_object_or_404(Propiedad,pk=pk)
    if propiedad:
        propiedad.delete()
        messages.success(request,"!Propiedad eliminada con exitó ")
        return redirect('listar_propiedades')
    
        
    
            
        